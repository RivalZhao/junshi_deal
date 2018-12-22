from lxml.etree import ElementTextIterator
from selenium import webdriver
from selenium.webdriver.common.action_chains import ActionChains
from time import sleep
import urllib
import docx
import docx.enum.style
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
import socket
import re
from lxml import etree
from lxml import html

pic_num = 0

def get_url(driver, url, text_xpath, button_xpath, url_xpath):
    driver.get(url)
    actions = ActionChains(driver)
    driver.implicitly_wait(5)
    driver.execute_script("window.scrollTo(0,1500)")
    driver.maximize_window()
    driver.implicitly_wait(5)
    sleep(3)
    driver.execute_script("window.scrollTo(0,5500)")
    driver.implicitly_wait(5)
    sleep(3)
    driver.execute_script("window.scrollTo(0,5500)")
    driver.implicitly_wait(5)
    sleep(1)
    driver.find_element_by_xpath(text_xpath).send_keys("40") #第几页 改这里~~
    sleep(1)
    click_node = driver.find_element_by_xpath(button_xpath)
    actions.click(click_node)
    actions.perform()

    driver.implicitly_wait(5)
    print('abc')
    sleep(3)
    driver.execute_script("window.scrollTo(0,1500)")
    print('abc11111')
    driver.implicitly_wait(5)
    sleep(3)
    driver.execute_script("window.scrollTo(0,5500)")
    print('abc22222')
    driver.implicitly_wait(5)
    sleep(4)
    driver.execute_script("window.scrollTo(0,6500)")
    print('abc33333')
    driver.implicitly_wait(5)
    sleep(3)
    nodes = driver.find_elements_by_xpath(url_xpath)
    url_list = []
    for node in nodes:
        url_list.append(node.get_attribute('href'))
        #print(node.get_attribute('href'))
    return url_list

def get_all_text(driver, url_list):
    content = []
    for i in range(len(url_list)):
        driver.get(url_list[i])
        driver.implicitly_wait(5)
        #标题 /html/body/div[4]/h1
        #//*[@id="top_bar"]/div/div[2]/span
        #//*[@id="top_bar"]/div/div[2]/a
        #//*[@id="article"]/p[1]
        #//*[@id="article"]/p[5]
        #//*[@id="article"]/div[2]/img
        #//*[@id="article"]/div[3]/img
        #广告 //*[@id="stage-obj11385866363843"]/a/img
        #结束//*[@id="article"]/p[11]/strong/text()

        #/html/body/div[4]/h1
        # node = driver.find_element_by_xpath('/html/body/div[3]')
        # for n in node.find_elements_by_xpath('child::h1'):
        #     print('cccccccccccc'+n.text)
        node = driver.find_element_by_class_name("main-title")
        print(node.text)
        content.append(node.text)
        node = driver.find_element_by_xpath('//*[@id="top_bar"]/div/div[2]/span')
        temp = node.text
        node = driver.find_element_by_xpath('//*[@id="top_bar"]/div/div[2]/a')
        temp += " " + node.text
        content.append(temp)
        node = driver.find_element_by_xpath('//*[@id="article"]')
        #for node in nodes:
        for child in node.find_elements_by_xpath('child::*'):
            if child.tag_name == 'p':
                if child.find_elements_by_xpath('child::strong'):
                    break
                else:
                    content.append(child.text)
            elif child.tag_name == 'div':
                if child.find_elements_by_xpath('child::img'):
                    global pic_num
                    pic_num = pic_num+1
                    imgurl = child.find_element_by_xpath('child::img').get_attribute('src')
                    point_index = (imgurl.rfind('.') + 1)
                    extension_name = imgurl[point_index: len(imgurl)]
                    save_picture(imgurl, pic_num, extension_name)
                    content.append("$$$$$$___$$$"+str(pic_num)+"."+extension_name+"$$$$$$___$$$")
            print('loop')
        print('222')
    export_word(content)

def get_all_text_x(url_list):
    print(url_list[0])
    print('in:get_all_text_x')
    content = []
    for i in range(len(url_list)):
        try:
            print(url_list[i])
            print('abc')
            sleep(1)
            print('start'+str(i))
            html_x = getHTML(url_list[i])
            #dom_tree = etree.HTML(html)
            dom_tree = html.fromstring(html_x)
            #标题 /html/body/div[4]/h1
            #//*[@id="top_bar"]/div/div[2]/span
            #//*[@id="top_bar"]/div/div[2]/a
            #//*[@id="article"]/p[1]
            #//*[@id="article"]/p[5]
            #//*[@id="article"]/div[2]/img
            #//*[@id="article"]/div[3]/img
            #广告 //*[@id="stage-obj11385866363843"]/a/img
            #结束//*[@id="article"]/p[11]/strong/text()

            #/html/body/div[4]/h1
            # node = driver.find_element_by_xpath('/html/body/div[3]')
            # for n in node.find_elements_by_xpath('child::h1'):
            #     print('cccccccccccc'+n.text)
            node = dom_tree.xpath('//*[@class="main-title"]')
            print(node[0].text)
            content.append(node[0].text)
            node = dom_tree.xpath('//*[@id="top_bar"]/div/div[2]/span')
            temp = node[0].text
            try:
                node = dom_tree.xpath('//*[@id="top_bar"]/div/div[2]/a')
                temp += " " + node[0].text
            except:
                node = dom_tree.xpath('//*[@id="top_bar"]/div/div[2]/span[2]')
                temp += " " + node[0].text
                print('except ttttttttttttttttttttt '+node[0].text)
            #print(temp)
            content.append(temp)
            node = dom_tree.xpath('//*[@id="article"]')
            #for node in nodes:
            for child in node[0]:
                #print(child.tag)
                if child.tag == 'p':
                    if child.text_content():
                        if child.text_content().find('本栏目所有文章目的在于传递更多信息，并不代表本网赞同其观点和对其真实性负责') != -1:
                            break
                        else:
                            content.append(child.text_content())
                            #print(child.text_content())
                            # for c in child: 这个只能取出里面的strong的文字
                            #     print(c.text)
                    # elif child.xpath('child::strong'):
                    #     if child.xpath('child::strong').text.find('本栏目所有文章目的在于传递更多信息，并不代表本网赞同其观点和对其真实性负责') != -1:
                    #         break
                    #     else:
                    #         content.append(child.xpath('child::strong').text)
                    #         print(child.xpath('child::strong').text)
                elif child.tag == 'div':
                    if child.xpath('child::img'):
                        global pic_num
                        pic_num = pic_num+1
                        imgurl = child.xpath('child::img')[0].attrib.get("src")
                        #print(imgurl)
                        point_index = (imgurl.rfind('.') + 1)
                        extension_name = imgurl[point_index: len(imgurl)]
                        save_picture(imgurl, pic_num, extension_name)
                        content.append("$$$$$$___$$$"+str(pic_num)+"."+extension_name+"$$$$$$___$$$")
                    if child.xpath('child::span'):
                        if child.xpath('child::span')[0].text_content():
                            content.append(child.xpath('child::span')[0].text_content())
                            print(child.xpath('child::span')[0].text_content())
                #print('loop')
        except:
            print('exceptexceptexceptexceptexceptexceptexceptexceptexceptexcept')
        print('222end')
    export_word(content)


def export_word(content):
    file = docx.Document()
    for con in content:
        match = re.search('(\$\$\$\$\$\$___\$\$\$)(.*?)(\$\$\$\$\$\$___\$\$\$)', con)
        if match:
            paragraph = file.add_picture('F:/junshi/'+match.group(2))
            # paragraph_format = paragraph.paragraph_format
            # paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        else:
            paragraph = file.add_paragraph(con)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph_format.line_spacing = Pt(11)
            #print(con)
    file.save("F:/2017.docx")

def getHTML(url):
    req_header = {
        'User-Agent': 'Mozilla/5.0 (Windows; U; Windows NT 6.1; en-US; rv:1.9.1.6) Gecko/20091201 Firefox/3.5.6'
    }
    req_timeout = 20
    html = 'cuowu'
    try:
        req = urllib.request.Request(url, None, req_header)
        resp = urllib.request.urlopen(req, None, req_timeout)
        html = resp.read()
    except urllib.error.URLError as e:
        print('cuowu la')
        print(e.message)
    except socket.timeout as e:
        print('timeout la')
        getHTML(url)
    return html

def save_picture(imgurl, num, extension_name):
    num = str(num)
    path = 'F:/junshi/' + num + '.' + extension_name
    content2 = urllib.request.urlopen(imgurl).read()
    with open(path, 'wb') as code:
        code.write(content2)




if __name__=="__main__":
    #driver = webdriver.PhantomJS(executable_path=r'E:/CloudStationBackup/PyCharm/KG_one/phantomjs-2.1.1-windows/bin/phantomjs')

    driver=webdriver.Firefox()
    # 获取指定页数上的网址列表
    # driver, url, text_xpath, button_xpath, url_xpath
    #/html/body/div[5]/div[4]/div[1]/div/div[4]/p/span[2]/input
    military_url_list = get_url(driver, 'http://mil.news.sina.com.cn/jssd/',
                         '/html/body/div[5]/div[4]/div[1]/div/div[4]/p/span[2]/input',
                         '/html/body/div[5]/div[4]/div[1]/div/div[4]/p/span[2]/a',
                         '/html/body/div[5]/div[4]/div[1]/div/div[2]/div/h3/a')
    # url_xpath

    #/html/body/div[4]/div[4]/div[1]/div/div[2]/div[2]/h3/a
    #/html/body/div[5]/div[4]/div[1]/div/div[2]/div[29]/h3/a

    #military_url_list = ['http://mil.news.sina.com.cn/jssd/2018-02-13/doc-ifyrmfmc2289510.shtml']
    #get_all_text(driver, military_url_list) 弃用 selenium速度慢
    get_all_text_x(military_url_list)


    #driver.close()